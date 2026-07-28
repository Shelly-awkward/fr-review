# -*- coding: utf-8 -*-
r"""
gen_checklist_docx.py — 管區意見內容 JSON → Word（python-docx 版）。

    python scripts/gen_checklist_docx.py out/8304_114_checklist_content.json out/管區意見.docx \
        [--cover out/複核表.docx]

內容 JSON 含 "cover"（財務報告實質審閱案件複核表）時，另存一份複核表 docx：
--cover 指定路徑；未指定則存於管區意見同目錄、檔名加「_複核表」。
與 gen_checklist_docx.js（docx-js 版）產出同體例（.js 版尚未支援複核表），擇一使用即可：
沒有 Node 的環境（例如只有 Python 沙箱的 AI）用本檔，其餘用 .js 版。
需求：pip install python-docx
"""
import argparse
import json
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

FONT = "PMingLiU"          # 新細明體，貼近原模板
W = {"g": 1300, "t": 5200, "y": 900, "n": 900, "b": 1740}   # DXA（＝twips）

# 需要人工處理的字樣→整句標紅字，避免承辦漏看
RED = RGBColor(0xC0, 0x00, 0x00)
RED_MARKERS = ("請自行", "洽公司", "請公司說明", "請向公司", "第＿頁", "查填")


def _segments(text):
    """切段供標紅：先把「（…第＿頁…）」括號段獨立出來（只紅括號、不紅整句），
    其餘依句號／分號切句。"""
    for part in re.split(r"(（[^（）]*第＿頁[^（）]*）)", text):
        if not part:
            continue
        if part.startswith("（") and "第＿頁" in part:
            yield part
            continue
        buf = ""
        for ch in part:
            buf += ch
            if ch in "。；":
                yield buf
                buf = ""
        if buf:
            yield buf


def styled(run, size=12, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    # 中文字型須另設 eastAsia，否則 Word 會用預設字型渲染中文
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return run


def add_text(p, text, size=12, bold=False):
    """寫入一段文字：含人工待辦字樣的句段用紅字。"""
    if not text:
        styled(p.add_run(""), size, bold)
        return
    for seg in _segments(text):
        run = styled(p.add_run(seg), size, bold)
        if any(m in seg for m in RED_MARKERS):
            run.font.color.rgb = RED


def para(container, text="", size=12, bold=False, align=None, indent=None,
         after=4, line=None):
    p = container.add_paragraph()
    add_text(p, text, size, bold)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    if indent is not None:
        pf.left_indent = Twips(indent)
    if line is not None:
        pf.line_spacing = line
    return p


def fill_cell(cell, lines, size=12, bold=False, align=None):
    """把多行文字寫進儲存格（清掉 python-docx 預設的空段落）。"""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell._element.clear_content()
    for line in (lines or [""]):
        p = cell.add_paragraph()
        add_text(p, line, size, bold)
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(0)


def build_checklist_table(doc, groups):
    n_items = sum(len(g["items"]) for g in groups)
    table = doc.add_table(rows=2 + n_items, cols=5)
    table.style = "Table Grid"
    table.autofit = False

    # 表頭兩列：內容／檢查內容（各跨兩列）＋填報項目（跨三欄）→ 是／否／備註
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    table.cell(0, 2).merge(table.cell(0, 4))
    fill_cell(table.cell(0, 0), ["內容"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(table.cell(0, 1), ["檢查內容"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(table.cell(0, 2), ["填報項目"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for col, (head, sub) in enumerate([("是", "(正常)"), ("否", "(異常)"), ("備註", None)], 2):
        cell = table.cell(1, col)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell._element.clear_content()
        for text, size in [(head, 12)] + ([(sub, 10)] if sub else []):
            p = cell.add_paragraph()
            styled(p.add_run(text), size, bold=(size == 12))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)

    r = 2
    for g in groups:
        first = r
        for it in g["items"]:
            fill_cell(table.cell(r, 1), it["text"].split("\n"), size=11)
            fill_cell(table.cell(r, 2), ["V" if it.get("mark") == "yes" else ""],
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            fill_cell(table.cell(r, 3), ["V" if it.get("mark") == "no" else ""],
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            fill_cell(table.cell(r, 4),
                      [t for t in (it.get("note") or "").split("\n") if t], size=10)
            r += 1
        merged = table.cell(first, 0)
        if r - 1 > first:
            merged = merged.merge(table.cell(r - 1, 0))
        fill_cell(merged, [g["group"]])

    for row in table.rows:
        for cell, key in zip(row.cells, ["g", "t", "y", "n", "b"]):
            cell.width = Twips(W[key])
    return table


def add_mini_table(doc, rows):
    """把連續的 【表】a｜b｜c 段落轉成小表格。"""
    n = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n)
    table.style = "Table Grid"
    for ri, cells in enumerate(rows):
        for ci, text in enumerate(cells):
            align = (WD_ALIGN_PARAGRAPH.CENTER
                     if ri == 0 or ci == n - 1 else None)
            fill_cell(table.cell(ri, ci), [text], size=10, bold=(ri == 0), align=align)
    para(doc, "", after=2)


def build_cover_doc(cover):
    """複核表（財務報告實質審閱案件複核表）——體例照過去實審樣本，單獨一份 docx。"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Twips(1080)
        section.left_margin = section.right_margin = Twips(1080)

    para(doc, "財務報告實質審閱案件複核表　106.10修正", size=15, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.alignment = 2  # 靠右
    for ri, (k, v) in enumerate([("保存期限", cover.get("保存期限", "")),
                                 ("檔　號", cover.get("檔號", ""))]):
        fill_cell(t.cell(ri, 0), [k], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(t.cell(ri, 1), [v], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        t.cell(ri, 0).width = Twips(1400)
        t.cell(ri, 1).width = Twips(1400)
    para(doc, "", after=4)

    para(doc, f"公司編號: {cover.get('公司編號', '')}　　"
              f"公司名稱: {cover.get('公司名稱', '')}", after=6)
    para(doc, f"財務報告年度期別︰{cover.get('財務報告年度期別', '')}", after=6)
    para(doc, f"公司背景介紹：{cover.get('公司背景介紹', '')}", after=8, line=1.3)

    para(doc, "公司之風險事項：", bold=True, after=3)
    para(doc, f"風險事項：{cover.get('風險事項', '')}", indent=360, after=3)
    para(doc, f"理由及因應措施：{cover.get('理由及因應措施', '')}", indent=360, after=8)

    para(doc, f"所屬產業本期成長（或衰退）趨勢：{cover.get('所屬產業趨勢', '')}", after=8)

    para(doc, "本次實質審閱加強查核重點︰", bold=True, after=3)
    para(doc, cover.get("加強查核重點", ""), indent=360, after=3)
    para(doc, f"複核意見︰{cover.get('複核意見1', '')}", indent=360, after=8)

    para(doc, "檢視異常事項", bold=True, after=3)
    for item in cover.get("檢視異常事項", []):
        para(doc, item, indent=360, after=3)
    para(doc, f"複核意見︰{cover.get('複核意見2', '')}", indent=360, after=8)

    para(doc, f"結論及擬辦：{cover.get('結論及擬辦', '')}", after=18, line=1.3)

    para(doc, "承辦　　　　科長　　　　複核　　　　副組長　　　　組長",
         align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    return doc


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("content")
    ap.add_argument("out")
    ap.add_argument("--cover", default="", help="複核表輸出路徑（內容 JSON 有 cover 才產出）")
    a = ap.parse_args()
    with open(a.content, encoding="utf-8") as f:
        c = json.load(f)

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Twips(1080)
        section.left_margin = section.right_margin = Twips(1080)

    para(doc, c["title"], size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    build_checklist_table(doc, c["groups"])
    for note in c.get("footnotes", []):
        para(doc, note, size=10, after=2)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for s in c["sections"]:
        para(doc, s["title"], size=13, bold=True, after=6)
        for b in s["body"]:
            if b.get("h"):
                para(doc, b["h"], bold=True, indent=360, after=3)
            pending = []
            for p in b.get("paras", []):
                if p.startswith("【表】"):
                    pending.append(p[3:].split("｜"))
                    continue
                if pending:
                    add_mini_table(doc, pending)
                    pending = []
                para(doc, p, indent=720 if b.get("h") else 480, after=4, line=1.3)
            if pending:
                add_mini_table(doc, pending)
        para(doc, "", after=3)

    doc.save(a.out)
    print(f"寫出 {a.out}")

    if c.get("cover"):
        cover_path = a.cover or (os.path.splitext(a.out)[0] + "_複核表.docx")
        build_cover_doc(c["cover"]).save(cover_path)
        print(f"寫出 {cover_path}")


if __name__ == "__main__":
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    main()
